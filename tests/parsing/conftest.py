"""Fixtures for the parsing stage: synthetic PDF / DOCX files."""
from __future__ import annotations

import zlib
import struct
from pathlib import Path
from typing import Callable

import fitz
import pytest
from docx import Document
from docx.shared import Pt


def _tiny_png(width: int = 8, height: int = 8) -> bytes:
    """Minimal grayscale PNG bytes."""

    def chunk(typ: bytes, data: bytes) -> bytes:
        c = typ + data
        return struct.pack(">I", len(data)) + c + struct.pack(
            ">I", zlib.crc32(c) & 0xFFFFFFFF
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 0, 0, 0, 0)
    raw = b"".join(b"\x00\xff\x00\xff\x00\xff\x00\xff" for _ in range(height))
    idat = zlib.compress(raw)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


@pytest.fixture
def make_pdf(tmp_dir: Path) -> Callable[..., Path]:
    """Factory that writes a PDF with configurable text/image pages."""

    def _build(
        pages: list[dict],
        name: str = "paper.pdf",
    ) -> Path:
        doc = fitz.open()
        for spec in pages:
            p = doc.new_page()
            text = spec.get("text", "")
            if spec.get("image"):
                p.insert_image(fitz.Rect(0, 0, 200, 200), stream=_tiny_png())
            if text:
                p.insert_text((72, 72), text)
        path = tmp_dir / name
        doc.save(str(path))
        doc.close()
        return path

    return _build


@pytest.fixture
def make_docx(tmp_dir: Path) -> Callable[..., Path]:
    """Factory that writes a DOCX with heading / plain paragraphs."""

    def _build(
        items: list[tuple[str, str]],
        name: str = "paper.docx",
    ) -> Path:
        doc = Document()
        for style, text in items:
            if style.startswith("Heading"):
                doc.add_heading(text, level=int(style.split()[-1]))
            elif style == "Title":
                doc.add_heading(text, level=0)
            else:
                doc.add_paragraph(text)
        path = tmp_dir / name
        doc.save(str(path))
        return path

    return _build
